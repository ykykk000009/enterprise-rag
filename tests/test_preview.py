from docx import Document as DocxDocument

from enterprise_document_rag.preview import render_document_preview


def test_docx_preview_renders_headings_tables_and_matching_context(tmp_path) -> None:
    source = tmp_path / "experiment.docx"
    document = DocxDocument()
    document.add_heading("3.2 Sample handling", level=1)
    document.add_paragraph("Solid mineral samples are eluted with PBS buffer.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Buffer"
    table.rows[0].cells[1].text = "PBS"
    document.save(source)

    rendered = render_document_preview(
        source_path=source,
        indexed_chunks=[
            {
                "id": "chunk-1",
                "text": "3.2 Sample handling Solid mineral samples are eluted with PBS buffer.",
                "page_no": None,
                "section_path": "3.2 Sample handling",
            }
        ],
        focus_chunk_id="chunk-1",
    )

    assert "<h1" in rendered
    assert "Solid mineral samples are eluted with PBS buffer." in rendered
    assert "<table>" in rendered
    assert 'id="matched-context"' in rendered

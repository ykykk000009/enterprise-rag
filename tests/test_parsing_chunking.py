from pathlib import Path
from unittest.mock import patch
from zipfile import ZipFile

import pymupdf
from docx import Document as DocxDocument
from openpyxl import Workbook

from enterprise_document_rag.chunking import StructureAwareChunker
from enterprise_document_rag.ocr import OcrTextBlock
from enterprise_document_rag.parsing import ParsedBlock, ParsedDocument, ParsedPage, parse_document


def test_pdf_parser_retains_page_and_bbox_metadata(tmp_path: Path) -> None:
    pdf_path = tmp_path / "sample.pdf"
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "Technical Parameters\nFlow rate 20 L/min")
    document.save(pdf_path)
    document.close()

    parsed = parse_document(pdf_path)

    assert parsed.pages[0].page_no == 1
    assert parsed.blocks[0].page_no == 1
    assert parsed.blocks[0].bbox is not None
    assert "Flow rate" in parsed.blocks[0].text


def test_pdf_uses_ocr_only_for_low_text_page(tmp_path: Path) -> None:
    class StubOcrProvider:
        def __init__(self) -> None:
            self.calls = 0

        def extract_page(self, *, page, dpi: int):
            del page
            self.calls += 1
            assert dpi == 150
            return (
                OcrTextBlock(
                    text="OCR extracted pressure is 0.25 MPa.",
                    bbox=(10.0, 20.0, 100.0, 40.0),
                    confidence=0.97,
                ),
            )

    pdf_path = tmp_path / "scanned.pdf"
    document = pymupdf.open()
    document.new_page()
    document.save(pdf_path)
    document.close()
    provider = StubOcrProvider()

    parsed = parse_document(
        pdf_path,
        ocr_enabled=True,
        ocr_min_text_chars_per_page=40,
        ocr_provider=provider,
    )

    assert provider.calls == 1
    assert parsed.blocks[0].block_type == "ocr"
    assert parsed.blocks[0].confidence == 0.97
    assert parsed.blocks[0].bbox == (10.0, 20.0, 100.0, 40.0)


def test_docx_parser_retains_section_metadata(tmp_path: Path) -> None:
    docx_path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.add_heading("3 Technical Parameters", level=1)
    document.add_paragraph("Cleaning pressure is configurable.")
    document.save(docx_path)

    parsed = parse_document(docx_path)

    paragraph = parsed.blocks[1]
    assert paragraph.section_path == ("3 Technical Parameters",)
    assert paragraph.text == "Cleaning pressure is configurable."


def test_excel_parser_repeats_sheet_headers_for_each_data_row(tmp_path: Path) -> None:
    workbook_path = tmp_path / "parameters.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "产品参数"
    worksheet.append(["型号", "流量", "压力"])
    worksheet.append(["PSG-1500", 1500, "0.25 MPa"])
    worksheet.append(["PSG-1800", 1800, "0.30 MPa"])
    workbook.save(workbook_path)

    parsed = parse_document(workbook_path)
    rows = [block for block in parsed.blocks if block.block_type == "table_row"]

    assert len(rows) == 2
    assert rows[0].section_path == ("工作表：产品参数",)
    assert "工作表：产品参数" in rows[0].text
    assert "行：2" in rows[0].text
    assert "型号：PSG-1500" in rows[0].text
    assert "流量：1500" in rows[0].text
    assert "表格标题" not in rows[0].text

    chunks = StructureAwareChunker(target_tokens=20, overlap_tokens=2, min_tokens=1).chunk(parsed)
    assert any("压力：0.30 MPa" in chunk.text for chunk in chunks)


def test_excel_parser_includes_cached_formula_values(tmp_path: Path) -> None:
    class Worksheet:
        title = "报价"

        def __init__(self, rows: list[tuple[object, ...]]) -> None:
            self.rows = rows

        def iter_rows(self, *, values_only: bool):
            assert values_only
            return iter(self.rows)

    class WorkbookStub:
        def __init__(self, rows: list[tuple[object, ...]]) -> None:
            self.worksheets = [Worksheet(rows)]

        def close(self) -> None:
            pass

    workbook_path = tmp_path / "quote.xlsx"
    workbook_path.touch()
    formula_workbook = WorkbookStub(
        [("报价明细",), ("项目", "医院总价"), ("合计", "=SUM(B2:B3)")]
    )
    cached_workbook = WorkbookStub(
        [("报价明细",), ("项目", "医院总价"), ("合计", 347.16)]
    )

    with patch(
        "enterprise_document_rag.parsing.load_workbook",
        side_effect=[formula_workbook, cached_workbook],
    ):
        parsed = parse_document(workbook_path)

    assert "医院总价：=SUM(B2:B3)（计算值：347.16）" in parsed.blocks[-1].text
    assert "表格标题：报价明细" in parsed.blocks[-1].text


def test_zip_parser_indexes_supported_member_paths_and_content(tmp_path: Path) -> None:
    workbook_path = tmp_path / "parameters.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "型号表"
    worksheet.append(["型号", "产量"])
    worksheet.append(["PSG-1500", 1500])
    workbook.save(workbook_path)

    archive_path = tmp_path / "tender-materials.zip"
    with ZipFile(archive_path, "w") as archive:
        archive.write(workbook_path, "报价/parameters.xlsx")
        archive.writestr("说明/readme.md", "# 范围\n\n仅用于内镜清洗项目。")

    parsed = parse_document(archive_path)
    chunks = StructureAwareChunker(target_tokens=20, overlap_tokens=2, min_tokens=1).chunk(parsed)
    combined = "\n".join(chunk.text for chunk in chunks)

    assert "压缩包内文件：报价/parameters.xlsx" in combined
    assert "型号：PSG-1500" in combined
    assert "压缩包内文件：说明/readme.md" in combined


def test_markdown_parser_and_chunker_preserve_structure_boundaries(tmp_path: Path) -> None:
    md_path = tmp_path / "sample.md"
    md_path.write_text(
        "# Product\n"
        "Alpha beta gamma.\n\n"
        "## Safety\n"
        "Use filtered water and record disinfection cycles.\n",
        encoding="utf-8",
    )

    parsed = parse_document(md_path)
    chunks = StructureAwareChunker(target_tokens=12, overlap_tokens=2, min_tokens=1).chunk(parsed)

    assert parsed.blocks[0].block_type == "heading"
    assert parsed.blocks[-1].section_path == ("Product", "Safety")
    assert len(chunks) >= 2
    assert chunks[0].chunk_index == 0
    assert chunks[0].next_chunk_index == 1


def test_large_text_is_chunked_without_dropping_content(tmp_path: Path) -> None:
    txt_path = tmp_path / "large.txt"
    paragraphs = [f"Paragraph {index} " + "word " * 80 for index in range(30)]
    txt_path.write_text("\n\n".join(paragraphs), encoding="utf-8")

    parsed = parse_document(txt_path)
    chunks = StructureAwareChunker(
        target_tokens=120,
        overlap_tokens=10,
        max_tokens=180,
    ).chunk(parsed)

    combined = "\n".join(chunk.text for chunk in chunks)
    assert "Paragraph 0" in combined
    assert "Paragraph 29" in combined
    assert len(chunks) > 1
    assert max(chunk.token_count for chunk in chunks) <= 260


def test_single_oversized_block_is_split_at_the_configured_maximum() -> None:
    text = "参数" * 500
    parsed = ParsedDocument(
        source_path="memory.txt",
        parser_version="test",
        pages=(
            ParsedPage(
                page_no=None,
                blocks=(ParsedBlock(text=text, page_no=None),),
            ),
        ),
    )

    chunks = StructureAwareChunker(
        target_tokens=50,
        overlap_tokens=5,
        min_tokens=1,
        max_tokens=80,
    ).chunk(parsed)

    assert len(chunks) > 1
    assert max(chunk.token_count for chunk in chunks) <= 80
    assert "".join(chunk.text for chunk in chunks) == text


def test_exact_duplicate_chunks_are_removed_within_a_document() -> None:
    parsed = ParsedDocument(
        source_path="memory.txt",
        parser_version="test",
        pages=(
            ParsedPage(
                page_no=None,
                blocks=(
                    ParsedBlock(text="重复内容", page_no=None),
                    ParsedBlock(text="重复内容", page_no=None),
                    ParsedBlock(text="唯一内容", page_no=None),
                ),
            ),
        ),
    )

    chunks = StructureAwareChunker(
        target_tokens=4,
        overlap_tokens=0,
        min_tokens=1,
        max_tokens=8,
    ).chunk(parsed)

    assert [chunk.text for chunk in chunks] == ["重复内容", "唯一内容"]
    assert [chunk.chunk_index for chunk in chunks] == [0, 1]

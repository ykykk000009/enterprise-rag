import gzip
import hashlib
import os
import re
import shutil
import subprocess
import tarfile
import tempfile
from collections import Counter
from collections.abc import Iterable
from dataclasses import dataclass, field, replace
from datetime import date, datetime
from io import BytesIO
from itertools import zip_longest
from pathlib import Path
from statistics import median
from typing import Any
from xml.etree import ElementTree
from zipfile import BadZipFile, ZipFile

import pymupdf
from docx import Document as DocxDocument
from openpyxl import load_workbook

from .ocr import OcrProvider, RapidOcrProvider
from .resource_control import background_work_gate
from .text_utils import sanitize_unicode
from .vision import ImageDescriptionProvider

PARSER_VERSION = "parser-v5"
PDF_LAYOUT_VERSION = "pdf-layout-v1"
ARCHIVE_MEMBER_SUFFIXES = frozenset(
    {".pdf", ".doc", ".docx", ".ppt", ".pptx", ".xlsx", ".xlsm", ".xls", ".txt", ".md"}
)


@dataclass(frozen=True)
class LayoutBlock:
    text: str
    page_no: int | None
    section_path: tuple[str, ...] = ()
    bbox: tuple[float, float, float, float] | None = None
    block_type: str = "paragraph"
    confidence: float | None = None
    font_size: float | None = None
    font_name: str | None = None
    is_bold: bool = False
    reading_order: int | None = None
    heading_level: int | None = None
    source_type: str = "native_text"
    table_markdown: str | None = None
    image_path: str | None = None
    caption: str | None = None
    image_metadata: dict[str, object] | None = None


@dataclass(frozen=True)
class ParsedBlock(LayoutBlock):
    """Backward-compatible name for blocks emitted by all document parsers."""


@dataclass(frozen=True)
class Section:
    title: str
    level: int
    parent_section: tuple[str, ...] | None
    children: tuple["Section", ...] = ()
    page_range: tuple[int, int] | None = None


@dataclass(frozen=True)
class ParsedPage:
    page_no: int | None
    blocks: tuple[LayoutBlock, ...]


@dataclass(frozen=True)
class ParsedDocument:
    source_path: str
    parser_version: str
    pages: tuple[ParsedPage, ...]
    document_structure_tree: tuple[Section, ...] = ()
    layout_version: str | None = None
    document_type: str | None = None

    @property
    def blocks(self) -> tuple[LayoutBlock, ...]:
        return tuple(block for page in self.pages for block in page.blocks)


@dataclass
class _SectionBuilder:
    title: str
    level: int
    path: tuple[str, ...]
    parent_path: tuple[str, ...] | None
    children: list["_SectionBuilder"] = field(default_factory=list)
    pages: set[int] = field(default_factory=set)


class UnsupportedFileTypeError(ValueError):
    pass


class ArchiveLimitError(ValueError):
    pass


class OfficeParsingError(ValueError):
    pass


def normalize_text(text: str) -> str:
    text = sanitize_unicode(text)
    lines = [" ".join(line.split()) for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def parse_document(
    path: str | Path,
    *,
    ocr_enabled: bool = False,
    ocr_min_text_chars_per_page: int = 40,
    ocr_render_dpi: int = 150,
    ocr_provider: OcrProvider | None = None,
    pdf_extract_images: bool = True,
    pdf_image_output_dir: str | Path | None = None,
    image_description_provider: ImageDescriptionProvider | None = None,
    archive_max_members: int = 500,
    archive_max_member_bytes: int = 50 * 1024 * 1024,
    archive_max_uncompressed_bytes: int = 200 * 1024 * 1024,
    archive_max_compression_ratio: int = 100,
) -> ParsedDocument:
    resolved = Path(path).resolve(strict=True)
    suffix = resolved.suffix.lower()
    if suffix == ".pdf":
        return parse_pdf(
            resolved,
            ocr_enabled=ocr_enabled,
            ocr_min_text_chars_per_page=ocr_min_text_chars_per_page,
            ocr_render_dpi=ocr_render_dpi,
            ocr_provider=ocr_provider,
            extract_images=pdf_extract_images,
            image_output_dir=pdf_image_output_dir,
            image_description_provider=image_description_provider,
        )
    if suffix == ".docx":
        return parse_docx(
            resolved,
            ocr_enabled=ocr_enabled,
            ocr_min_text_chars=ocr_min_text_chars_per_page,
            ocr_provider=ocr_provider,
            image_output_dir=pdf_image_output_dir,
            image_description_provider=image_description_provider,
        )
    if suffix == ".doc":
        return parse_doc(
            resolved,
            ocr_enabled=ocr_enabled,
            ocr_min_text_chars=ocr_min_text_chars_per_page,
            ocr_provider=ocr_provider,
            image_output_dir=pdf_image_output_dir,
            image_description_provider=image_description_provider,
        )
    if suffix in {".ppt", ".pptx"}:
        return parse_presentation(resolved)
    if suffix in {".xlsx", ".xlsm"}:
        return parse_xlsx(resolved)
    if suffix == ".xls":
        return parse_xls(resolved)
    if suffix == ".zip":
        return parse_zip(
            resolved,
            ocr_enabled=ocr_enabled,
            ocr_min_text_chars_per_page=ocr_min_text_chars_per_page,
            ocr_render_dpi=ocr_render_dpi,
            ocr_provider=ocr_provider,
            max_members=archive_max_members,
            max_member_bytes=archive_max_member_bytes,
            max_uncompressed_bytes=archive_max_uncompressed_bytes,
            max_compression_ratio=archive_max_compression_ratio,
        )
    if suffix in {".tar", ".gz"}:
        return parse_tar_or_gzip(
            resolved,
            ocr_enabled=ocr_enabled,
            ocr_min_text_chars_per_page=ocr_min_text_chars_per_page,
            ocr_render_dpi=ocr_render_dpi,
            ocr_provider=ocr_provider,
            max_members=archive_max_members,
            max_member_bytes=archive_max_member_bytes,
            max_uncompressed_bytes=archive_max_uncompressed_bytes,
            max_compression_ratio=archive_max_compression_ratio,
        )
    if suffix in {".rar", ".7z"}:
        return parse_external_archive(
            resolved,
            ocr_enabled=ocr_enabled,
            ocr_min_text_chars_per_page=ocr_min_text_chars_per_page,
            ocr_render_dpi=ocr_render_dpi,
            ocr_provider=ocr_provider,
            max_members=archive_max_members,
            max_member_bytes=archive_max_member_bytes,
            max_uncompressed_bytes=archive_max_uncompressed_bytes,
            max_compression_ratio=archive_max_compression_ratio,
        )
    if suffix == ".txt":
        return parse_text_file(resolved, markdown=False)
    if suffix == ".md":
        return parse_text_file(resolved, markdown=True)
    raise UnsupportedFileTypeError(f"unsupported file type: {suffix}")


def parse_pdf(
    path: str | Path,
    *,
    ocr_enabled: bool = False,
    ocr_min_text_chars_per_page: int = 40,
    ocr_render_dpi: int = 150,
    ocr_provider: OcrProvider | None = None,
    extract_images: bool = True,
    image_output_dir: str | Path | None = None,
    image_description_provider: ImageDescriptionProvider | None = None,
) -> ParsedDocument:
    resolved = Path(path).resolve(strict=True)
    with pymupdf.open(resolved) as document:
        return _parse_pdf_document(
            document=document,
            source_path=str(resolved),
            ocr_enabled=ocr_enabled,
            ocr_min_text_chars_per_page=ocr_min_text_chars_per_page,
            ocr_render_dpi=ocr_render_dpi,
            ocr_provider=ocr_provider,
            extract_images=extract_images,
            image_output_dir=Path(image_output_dir) if image_output_dir is not None else None,
            image_description_provider=image_description_provider,
        )


def _parse_pdf_document(
    *,
    document,
    source_path: str,
    ocr_enabled: bool,
    ocr_min_text_chars_per_page: int,
    ocr_render_dpi: int,
    ocr_provider: OcrProvider | None,
    extract_images: bool = True,
    image_output_dir: Path | None = None,
    image_description_provider: ImageDescriptionProvider | None = None,
) -> ParsedDocument:
    page_blocks: list[tuple[int, float, float, list[ParsedBlock]]] = []
    provider = ocr_provider
    for page_index, page in enumerate(document, start=1):
        background_work_gate.wait_for_background_work()
        table_blocks = _extract_pdf_table_blocks(page=page, page_no=page_index)
        blocks = _extract_pdf_layout_blocks(page=page, page_no=page_index)
        blocks = _exclude_table_text_blocks(blocks, table_blocks=table_blocks)
        blocks.extend(table_blocks)
        image_regions = _pdf_image_regions(page=page, table_blocks=table_blocks)
        if ocr_enabled and (
            image_regions
            or _should_ocr_page(
                blocks=blocks,
                min_text_chars=ocr_min_text_chars_per_page,
            )
        ):
            if provider is None:
                provider = RapidOcrProvider()
            blocks.extend(
                _extract_pdf_ocr_blocks(
                    page=page,
                    page_no=page_index,
                    native_blocks=blocks,
                    image_regions=image_regions,
                    provider=provider,
                    dpi=ocr_render_dpi,
                    min_text_chars=ocr_min_text_chars_per_page,
                )
            )
        if extract_images and image_regions:
            image_blocks = _extract_pdf_image_blocks(
                page=page,
                page_no=page_index,
                source_path=source_path,
                regions=image_regions,
                context_blocks=blocks,
                output_dir=image_output_dir,
                description_provider=image_description_provider,
            )
            table_regions = _ocr_table_regions(
                regions=image_regions,
                image_blocks=image_blocks,
            )
            if table_regions:
                blocks = _exclude_ocr_table_cells(blocks, regions=table_regions)
            blocks.extend(image_blocks)
        ordered = _order_pdf_blocks(blocks, page_width=float(page.rect.width))
        page_blocks.append(
            (page_index, float(page.rect.width), float(page.rect.height), ordered)
        )

    classified_pages = _classify_pdf_layout(page_blocks)
    pages, section_tree = _reconstruct_pdf_sections(classified_pages)
    return ParsedDocument(
        source_path=source_path,
        parser_version=PARSER_VERSION,
        pages=tuple(pages),
        document_structure_tree=section_tree,
        layout_version=PDF_LAYOUT_VERSION,
        document_type=_detect_pdf_type(pages),
    )


def _pdf_image_regions(
    *,
    page,
    table_blocks: list[ParsedBlock],
) -> tuple[tuple[float, float, float, float], ...]:
    try:
        image_info = page.get_image_info(xrefs=True)
    except (AttributeError, RuntimeError, TypeError, ValueError):
        image_info = ()
    table_boxes = [block.bbox for block in table_blocks if block.bbox is not None]
    regions: list[tuple[float, float, float, float]] = []
    for info in image_info:
        bbox = _float_bbox(info.get("bbox") if isinstance(info, dict) else None)
        if bbox is None:
            continue
        if bbox[2] - bbox[0] < 8 or bbox[3] - bbox[1] < 8:
            continue
        if any(_bbox_overlap_ratio(bbox, table_bbox) >= 0.5 for table_bbox in table_boxes):
            continue
        regions.append(bbox)
    return _merge_image_regions(regions)


def _merge_image_regions(
    regions: list[tuple[float, float, float, float]],
) -> tuple[tuple[float, float, float, float], ...]:
    merged: list[tuple[float, float, float, float]] = []
    for region in sorted(regions, key=lambda item: (item[1], item[0])):
        for index, existing in enumerate(merged):
            intersection = _bbox_intersection_area(existing, region)
            smaller = min(_bbox_area(existing), _bbox_area(region))
            if smaller > 0 and intersection / smaller >= 0.8:
                merged[index] = _union_bboxes((existing, region)) or existing
                break
        else:
            merged.append(region)
    return tuple(merged)


def _extract_pdf_ocr_blocks(
    *,
    page,
    page_no: int,
    native_blocks: list[ParsedBlock],
    image_regions: tuple[tuple[float, float, float, float], ...],
    provider: OcrProvider,
    dpi: int,
    min_text_chars: int,
) -> list[ParsedBlock]:
    extracted = []
    extract_region = getattr(provider, "extract_region", None)
    if image_regions and callable(extract_region):
        for region in image_regions:
            try:
                region_blocks = extract_region(page=page, bbox=region, dpi=dpi)
            except (OSError, RuntimeError, TypeError, ValueError):
                continue
            extracted.extend(
                _ocr_layout_block(item, page_no=page_no, fallback_bbox=region)
                for item in region_blocks
                if normalize_text(item.text)
            )
    elif image_regions:
        extracted.extend(
            _ocr_layout_block(item, page_no=page_no)
            for item in provider.extract_page(page=page, dpi=dpi)
            if normalize_text(item.text)
        )

    if not extracted and _should_ocr_page(
        blocks=native_blocks,
        min_text_chars=min_text_chars,
    ):
        extracted.extend(
            _ocr_layout_block(item, page_no=page_no)
            for item in provider.extract_page(page=page, dpi=dpi)
            if normalize_text(item.text)
        )
    return _deduplicate_ocr_blocks(extracted, native_blocks=native_blocks)


def _ocr_layout_block(
    item,
    *,
    page_no: int,
    fallback_bbox: tuple[float, float, float, float] | None = None,
) -> ParsedBlock:
    return ParsedBlock(
        text=normalize_text(item.text),
        page_no=page_no,
        bbox=item.bbox or fallback_bbox,
        block_type="ocr",
        confidence=item.confidence,
        source_type="ocr_text",
    )


def _deduplicate_ocr_blocks(
    blocks: list[ParsedBlock],
    *,
    native_blocks: list[ParsedBlock],
) -> list[ParsedBlock]:
    native_keys = {_recurring_text_key(block.text) for block in native_blocks}
    seen: set[tuple[str, tuple[float, float, float, float] | None]] = set()
    unique = []
    for block in blocks:
        text_key = _recurring_text_key(block.text)
        if text_key in native_keys:
            continue
        bbox_key = (
            tuple(round(value, 1) for value in block.bbox)
            if block.bbox is not None
            else None
        )
        key = (text_key, bbox_key)
        if key not in seen:
            seen.add(key)
            unique.append(block)
    return unique


def _detect_pdf_type(pages: list[ParsedPage]) -> str:
    has_native = any(
        block.source_type == "native_text" and block.block_type not in {"header", "footer"}
        for page in pages
        for block in page.blocks
    )
    has_ocr = any(
        block.source_type == "ocr_text"
        for page in pages
        for block in page.blocks
    )
    if has_native and has_ocr:
        return "mixed_pdf"
    if has_ocr:
        return "scanned_pdf"
    return "native_pdf"


def _extract_pdf_image_blocks(
    *,
    page,
    page_no: int,
    source_path: str,
    regions: tuple[tuple[float, float, float, float], ...],
    context_blocks: list[ParsedBlock],
    output_dir: Path | None,
    description_provider: ImageDescriptionProvider | None,
) -> list[ParsedBlock]:
    image_blocks: list[ParsedBlock] = []
    for region in regions:
        try:
            pixmap = page.get_pixmap(
                matrix=pymupdf.Matrix(2, 2),
                clip=pymupdf.Rect(region),
                alpha=False,
            )
            image_bytes = pixmap.tobytes("png")
        except (OSError, RuntimeError, TypeError, ValueError):
            continue
        if not image_bytes:
            continue
        image_path = _save_pdf_image(
            image_bytes=image_bytes,
            source_path=source_path,
            page_no=page_no,
            output_dir=output_dir,
        )
        ocr_tables = _reconstruct_ocr_tables(
            region=region,
            page_no=page_no,
            context_blocks=context_blocks,
            image_path=str(image_path),
        )
        if ocr_tables:
            image_blocks.extend(ocr_tables)
            continue
        nearby_text = _nearby_image_caption(region=region, blocks=context_blocks)
        ocr_text = "\n".join(
            block.text
            for block in context_blocks
            if block.source_type == "ocr_text"
            and block.bbox is not None
            and _bbox_overlap_ratio(block.bbox, region) >= 0.5
        )
        caption = _describe_pdf_image(
            image_bytes=image_bytes,
            page_no=page_no,
            bbox=region,
            nearby_text=nearby_text,
            ocr_text=ocr_text,
            provider=description_provider,
        )
        image_blocks.append(
            ParsedBlock(
                text=caption,
                page_no=page_no,
                bbox=region,
                block_type="image",
                source_type="image_description",
                image_path=str(image_path),
                caption=caption,
            )
        )
    return image_blocks


def _reconstruct_ocr_tables(
    *,
    region: tuple[float, float, float, float],
    page_no: int,
    context_blocks: list[ParsedBlock],
    image_path: str,
) -> list[ParsedBlock]:
    """Turn OCR cells from an embedded image into conservative Markdown tables.

    PDF table extraction only sees native text/vector lines.  For a raster table,
    RapidOCR gives us one text item per cell together with its PDF coordinates.
    Reconstructing rows and stable columns here retains the relationships that a
    newline-only ``Image text`` caption loses.
    """
    ocr_cells = [
        block
        for block in context_blocks
        if block.source_type == "ocr_text"
        and block.bbox is not None
        and _bbox_overlap_ratio(block.bbox, region) >= 0.5
    ]
    row_groups = _group_ocr_cells_into_rows(ocr_cells)
    tables: list[ParsedBlock] = []
    for rows in _split_ocr_table_rows(row_groups):
        reconstructed = _ocr_rows_to_table(rows)
        if reconstructed is None:
            continue
        markdown, table_bbox, confidence = reconstructed
        tables.append(
            ParsedBlock(
                text=markdown,
                page_no=page_no,
                bbox=table_bbox,
                block_type="table",
                confidence=confidence,
                source_type="ocr_table",
                table_markdown=markdown,
                image_path=image_path,
                caption="OCR reconstructed table",
            )
        )
    return tables


def _ocr_table_regions(
    *,
    regions: tuple[tuple[float, float, float, float], ...],
    image_blocks: list[ParsedBlock],
) -> tuple[tuple[float, float, float, float], ...]:
    table_boxes = [
        block.bbox
        for block in image_blocks
        if block.source_type == "ocr_table" and block.bbox is not None
    ]
    return tuple(
        region
        for region in regions
        if any(_bbox_overlap_ratio(table_box, region) >= 0.8 for table_box in table_boxes)
    )


def _exclude_ocr_table_cells(
    blocks: list[ParsedBlock],
    *,
    regions: tuple[tuple[float, float, float, float], ...],
) -> list[ParsedBlock]:
    """Do not index both a reconstructed table and its lossy OCR cell stream."""
    return [
        block
        for block in blocks
        if block.source_type != "ocr_text"
        or block.bbox is None
        or not any(_bbox_overlap_ratio(block.bbox, region) >= 0.5 for region in regions)
    ]


def _group_ocr_cells_into_rows(cells: list[ParsedBlock]) -> list[list[ParsedBlock]]:
    if not cells:
        return []
    heights = [block.bbox[3] - block.bbox[1] for block in cells if block.bbox]
    row_tolerance = max(5.0, median(heights) * 0.75) if heights else 8.0
    rows: list[list[ParsedBlock]] = []
    for cell in sorted(cells, key=lambda block: (_bbox_center_y(block), _bbox_center_x(block))):
        for row in rows:
            row_center = median(_bbox_center_y(item) for item in row)
            if abs(_bbox_center_y(cell) - row_center) <= row_tolerance:
                row.append(cell)
                break
        else:
            rows.append([cell])
    return [sorted(row, key=_bbox_center_x) for row in rows]


def _split_ocr_table_rows(rows: list[list[ParsedBlock]]) -> list[list[list[ParsedBlock]]]:
    if len(rows) < 2:
        return []
    centers = [median(_bbox_center_y(cell) for cell in row) for row in rows]
    gaps = [
        later - earlier
        for earlier, later in zip(centers, centers[1:], strict=False)
    ]
    cell_heights = [
        cell.bbox[3] - cell.bbox[1]
        for row in rows
        for cell in row
        if cell.bbox is not None
    ]
    if not gaps or not cell_heights:
        return [rows]
    split_gap = max(median(cell_heights) * 2.4, median(gaps) * 2.0)
    groups: list[list[list[ParsedBlock]]] = [[rows[0]]]
    for gap, row in zip(gaps, rows[1:], strict=False):
        if gap > split_gap:
            groups.append([row])
        else:
            groups[-1].append(row)
    return groups


def _ocr_rows_to_table(
    rows: list[list[ParsedBlock]],
) -> tuple[str, tuple[float, float, float, float], float | None] | None:
    if len(rows) < 2:
        return None
    cells = [cell for row in rows for cell in row if cell.bbox is not None]
    if len(cells) < 4:
        return None
    widths = [cell.bbox[2] - cell.bbox[0] for cell in cells if cell.bbox]
    column_tolerance = max(12.0, min(36.0, median(widths) * 0.9)) if widths else 18.0
    columns: list[list[ParsedBlock]] = []
    for cell in sorted(cells, key=_bbox_center_x):
        for column in columns:
            column_center = median(_bbox_center_x(item) for item in column)
            if abs(_bbox_center_x(cell) - column_center) <= column_tolerance:
                column.append(cell)
                break
        else:
            columns.append([cell])

    minimum_column_coverage = max(2, (len(rows) * 3 + 4) // 5)
    stable_columns = [
        column
        for column in columns
        if sum(
            any(cell in row for cell in column)
            for row in rows
        ) >= minimum_column_coverage
    ]
    if len(stable_columns) < 2:
        return None
    centers = [median(_bbox_center_x(cell) for cell in column) for column in stable_columns]
    assignment_tolerance = column_tolerance * 1.25
    normalized_rows: list[tuple[str, ...]] = []
    selected_cells: list[ParsedBlock] = []
    dense_rows = 0
    for row in rows:
        values: list[str] = []
        for center in centers:
            candidates = [
                cell
                for cell in row
                if abs(_bbox_center_x(cell) - center) <= assignment_tolerance
            ]
            if candidates:
                chosen = max(candidates, key=lambda cell: cell.confidence or 0.0)
                values.append(chosen.text)
                selected_cells.append(chosen)
            else:
                values.append("")
        if sum(bool(value) for value in values) >= max(2, len(centers) - 1):
            dense_rows += 1
        normalized_rows.append(tuple(values))
    normalized = _normalize_table_rows(normalized_rows)
    if not _is_usable_table(normalized) or dense_rows < max(2, (len(rows) * 3 + 3) // 4):
        return None
    markdown = _table_to_markdown(normalized)
    table_bbox = _union_bboxes(
        cell.bbox for cell in selected_cells if cell.bbox is not None
    )
    if not markdown or table_bbox is None:
        return None
    confidences = [cell.confidence for cell in selected_cells if cell.confidence is not None]
    confidence = sum(confidences) / len(confidences) if confidences else None
    return markdown, table_bbox, confidence


def _bbox_center_x(block: ParsedBlock) -> float:
    assert block.bbox is not None
    return (block.bbox[0] + block.bbox[2]) / 2


def _bbox_center_y(block: ParsedBlock) -> float:
    assert block.bbox is not None
    return (block.bbox[1] + block.bbox[3]) / 2


def _save_pdf_image(
    *,
    image_bytes: bytes,
    source_path: str,
    page_no: int,
    output_dir: Path | None,
) -> Path:
    root = output_dir or Path(tempfile.gettempdir()) / "enterprise-document-rag-images"
    document_key = hashlib.sha256(source_path.encode("utf-8")).hexdigest()[:16]
    image_key = hashlib.sha256(image_bytes).hexdigest()[:16]
    directory = root.expanduser().resolve() / document_key
    directory.mkdir(parents=True, exist_ok=True)
    target = directory / f"page-{page_no:04d}-{image_key}.png"
    if not target.exists():
        target.write_bytes(image_bytes)
    return target


def _nearby_image_caption(
    *,
    region: tuple[float, float, float, float],
    blocks: list[ParsedBlock],
) -> str:
    candidates: list[tuple[float, str]] = []
    for block in blocks:
        if block.bbox is None or block.block_type in {"ocr", "image", "table"}:
            continue
        horizontal_overlap = max(
            0.0,
            min(region[2], block.bbox[2]) - max(region[0], block.bbox[0]),
        )
        if horizontal_overlap <= 0:
            continue
        if block.bbox[1] >= region[3]:
            distance = block.bbox[1] - region[3]
        elif region[1] >= block.bbox[3]:
            distance = region[1] - block.bbox[3]
        else:
            distance = 0.0
        if distance > 90:
            continue
        caption_bonus = 100.0 if re.match(
            r"^\s*(?:图\s*\d+|Figure\s*\d+|Fig\.\s*\d+)",
            block.text,
            re.IGNORECASE,
        ) else 0.0
        candidates.append((caption_bonus - distance, block.text))
    return max(candidates, default=(0.0, ""), key=lambda item: item[0])[1]


def _describe_pdf_image(
    *,
    image_bytes: bytes,
    page_no: int,
    bbox: tuple[float, float, float, float],
    nearby_text: str,
    ocr_text: str,
    provider: ImageDescriptionProvider | None,
) -> str:
    if provider is not None:
        try:
            description = normalize_text(
                provider.describe_image(
                    image_bytes=image_bytes,
                    page_no=page_no,
                    bbox=bbox,
                    nearby_text=nearby_text,
                    ocr_text=ocr_text,
                )
            )
        except (OSError, RuntimeError, TypeError, ValueError):
            description = ""
        if description:
            return description
    parts = []
    if nearby_text:
        parts.append(nearby_text)
    if ocr_text and ocr_text not in nearby_text:
        parts.append(f"Image text: {ocr_text}")
    if not parts:
        width = round(bbox[2] - bbox[0])
        height = round(bbox[3] - bbox[1])
        parts.append(f"Embedded image on PDF page {page_no} ({width} x {height} points)")
    return "\n".join(parts)


def _extract_pdf_layout_blocks(*, page, page_no: int) -> list[ParsedBlock]:
    """Extract text spans without flattening font and geometry information."""
    blocks: list[ParsedBlock] = []
    page_dict = page.get_text("dict", sort=False)
    for raw_block in page_dict.get("blocks", []):
        if raw_block.get("type", 0) != 0:
            continue
        lines = raw_block.get("lines", [])
        spans = [span for line in lines for span in line.get("spans", [])]
        line_texts = [
            normalize_text("".join(str(span.get("text", "")) for span in line.get("spans", [])))
            for line in lines
        ]
        text = "\n".join(item for item in line_texts if item)
        if not text:
            continue
        bbox = _float_bbox(raw_block.get("bbox"))
        if bbox is None:
            span_boxes = [
                item
                for span in spans
                if (item := _float_bbox(span.get("bbox"))) is not None
            ]
            bbox = _union_bboxes(span_boxes)
        font_weights: Counter[str] = Counter()
        bold_characters = 0
        total_characters = 0
        font_sizes: list[float] = []
        for span in spans:
            span_text = str(span.get("text", ""))
            weight = max(len(span_text.strip()), 1)
            font_name = str(span.get("font", "")).strip()
            if font_name:
                font_weights[font_name] += weight
            size = span.get("size")
            if isinstance(size, (int, float)) and size > 0:
                font_sizes.extend([float(size)] * min(weight, 200))
            total_characters += weight
            if _span_is_bold(span):
                bold_characters += weight
        blocks.append(
            ParsedBlock(
                text=text,
                page_no=page_no,
                bbox=bbox,
                block_type="list" if _looks_like_list_item(text) else "paragraph",
                font_size=max(font_sizes) if font_sizes else None,
                font_name=font_weights.most_common(1)[0][0] if font_weights else None,
                is_bold=total_characters > 0 and bold_characters / total_characters >= 0.5,
            )
        )
    return blocks


def _extract_pdf_table_blocks(*, page, page_no: int) -> list[ParsedBlock]:
    """Extract native PDF tables and retain their exact page coordinates."""
    extracted: list[ParsedBlock] = []
    seen: set[tuple[float, float, float, float]] = set()
    attempts = (
        {},
        {
            "vertical_strategy": "text",
            "horizontal_strategy": "text",
            "min_words_vertical": 2,
            "min_words_horizontal": 1,
        },
    )
    for options in attempts:
        try:
            finder = page.find_tables(**options)
        except (AttributeError, RuntimeError, TypeError, ValueError):
            continue
        for table in getattr(finder, "tables", ()):
            bbox = _float_bbox(getattr(table, "bbox", None))
            if bbox is None:
                continue
            bbox_key = tuple(round(value, 2) for value in bbox)
            if bbox_key in seen:
                continue
            try:
                rows = table.extract()
            except (AttributeError, RuntimeError, TypeError, ValueError):
                continue
            normalized_rows = _normalize_table_rows(rows)
            if not _is_usable_table(normalized_rows):
                continue
            if options and len(normalized_rows[0]) < 3:
                # Text-only inference is intentionally conservative: two-column
                # prose is otherwise frequently mistaken for a borderless table.
                continue
            markdown = _table_to_markdown(normalized_rows)
            if not markdown:
                continue
            seen.add(bbox_key)
            extracted.append(
                ParsedBlock(
                    text=markdown,
                    page_no=page_no,
                    bbox=bbox,
                    block_type="table",
                    table_markdown=markdown,
                )
            )
        if extracted:
            break
    return extracted


def _normalize_table_rows(rows: Any) -> tuple[tuple[str, ...], ...]:
    if not isinstance(rows, (list, tuple)):
        return ()
    normalized = [
        tuple(normalize_text(str(cell)) if cell is not None else "" for cell in row)
        for row in rows
        if isinstance(row, (list, tuple))
    ]
    normalized = [row for row in normalized if any(row)]
    if not normalized:
        return ()
    width = max(len(row) for row in normalized)
    padded = [(*row, *("" for _ in range(width - len(row)))) for row in normalized]
    populated_columns = [
        index
        for index in range(width)
        if any(row[index] for row in padded)
    ]
    return tuple(
        tuple(row[index] for index in populated_columns)
        for row in padded
    )


def _is_usable_table(rows: tuple[tuple[str, ...], ...]) -> bool:
    if len(rows) < 2 or not rows or len(rows[0]) < 2:
        return False
    populated = sum(bool(cell) for row in rows for cell in row)
    dense_rows = sum(sum(bool(cell) for cell in row) >= 2 for row in rows)
    return populated >= 4 and dense_rows >= max(2, len(rows) // 2)


def _table_to_markdown(rows: tuple[tuple[str, ...], ...]) -> str:
    if not rows:
        return ""
    width = len(rows[0])
    header = [
        _escape_markdown_cell(value) or f"Column {index + 1}"
        for index, value in enumerate(rows[0])
    ]
    body = [
        [_escape_markdown_cell(row[index]) for index in range(width)]
        for row in rows[1:]
    ]
    lines = [
        f"| {' | '.join(header)} |",
        f"| {' | '.join('---' for _ in range(width))} |",
    ]
    lines.extend(f"| {' | '.join(row)} |" for row in body)
    return "\n".join(lines)


def _escape_markdown_cell(value: str) -> str:
    return normalize_text(value).replace("|", r"\|").replace("\n", "<br>")


def _exclude_table_text_blocks(
    blocks: list[ParsedBlock],
    *,
    table_blocks: list[ParsedBlock],
) -> list[ParsedBlock]:
    table_boxes = [block.bbox for block in table_blocks if block.bbox is not None]
    if not table_boxes:
        return blocks
    return [
        block
        for block in blocks
        if block.bbox is None
        or not any(_bbox_overlap_ratio(block.bbox, table_bbox) >= 0.5 for table_bbox in table_boxes)
    ]


def _bbox_overlap_ratio(
    first: tuple[float, float, float, float],
    second: tuple[float, float, float, float],
) -> float:
    x0 = max(first[0], second[0])
    y0 = max(first[1], second[1])
    x1 = min(first[2], second[2])
    y1 = min(first[3], second[3])
    if x1 <= x0 or y1 <= y0:
        return 0.0
    intersection = (x1 - x0) * (y1 - y0)
    first_area = max((first[2] - first[0]) * (first[3] - first[1]), 1.0)
    return intersection / first_area


def _bbox_intersection_area(
    first: tuple[float, float, float, float],
    second: tuple[float, float, float, float],
) -> float:
    return max(0.0, min(first[2], second[2]) - max(first[0], second[0])) * max(
        0.0,
        min(first[3], second[3]) - max(first[1], second[1]),
    )


def _bbox_area(bbox: tuple[float, float, float, float]) -> float:
    return max(0.0, bbox[2] - bbox[0]) * max(0.0, bbox[3] - bbox[1])


def _span_is_bold(span: dict[str, Any]) -> bool:
    font_name = str(span.get("font", "")).casefold()
    flags = int(span.get("flags", 0) or 0)
    return bool(flags & pymupdf.TEXT_FONT_BOLD) or any(
        marker in font_name for marker in ("bold", "black", "heavy", "demi", "semibold")
    )


def _float_bbox(value: Any) -> tuple[float, float, float, float] | None:
    if all(hasattr(value, attribute) for attribute in ("x0", "y0", "x1", "y1")):
        value = (value.x0, value.y0, value.x1, value.y1)
    if not isinstance(value, (list, tuple)) or len(value) != 4:
        return None
    try:
        x0, y0, x1, y1 = (float(item) for item in value)
    except (TypeError, ValueError):
        return None
    if x1 <= x0 or y1 <= y0:
        return None
    return (x0, y0, x1, y1)


def _union_bboxes(
    boxes: Iterable[tuple[float, float, float, float]],
) -> tuple[float, float, float, float] | None:
    values = tuple(boxes)
    if not values:
        return None
    return (
        min(item[0] for item in values),
        min(item[1] for item in values),
        max(item[2] for item in values),
        max(item[3] for item in values),
    )


def _order_pdf_blocks(
    blocks: list[ParsedBlock],
    *,
    page_width: float,
) -> list[ParsedBlock]:
    """Produce stable reading order, including common two-column documents."""
    positioned = [block for block in blocks if block.bbox is not None]
    unpositioned = [block for block in blocks if block.bbox is None]
    if not positioned:
        ordered = unpositioned
    else:
        spanning = [
            block
            for block in positioned
            if block.bbox is not None and (block.bbox[2] - block.bbox[0]) >= page_width * 0.65
        ]
        regular = [block for block in positioned if block not in spanning]
        separators = sorted(spanning, key=_block_position_key)
        ordered = []
        lower_y = float("-inf")
        for separator in [*separators, None]:
            upper_y = (
                separator.bbox[1]
                if separator is not None and separator.bbox
                else float("inf")
            )
            band = [
                block
                for block in regular
                if block.bbox is not None
                and lower_y <= (block.bbox[1] + block.bbox[3]) / 2 < upper_y
            ]
            ordered.extend(_order_pdf_band(band, page_width=page_width))
            if separator is not None:
                ordered.append(separator)
                lower_y = separator.bbox[3] if separator.bbox else upper_y
        ordered.extend(unpositioned)
    return [replace(block, reading_order=index) for index, block in enumerate(ordered)]


def _order_pdf_band(blocks: list[ParsedBlock], *, page_width: float) -> list[ParsedBlock]:
    if len(blocks) < 4:
        return sorted(blocks, key=_block_position_key)
    midpoint = page_width / 2
    tolerance = page_width * 0.04
    left = [
        block
        for block in blocks
        if block.bbox is not None and block.bbox[2] <= midpoint + tolerance
    ]
    right = [
        block
        for block in blocks
        if block.bbox is not None and block.bbox[0] >= midpoint - tolerance
    ]
    if len(left) >= 2 and len(right) >= 2 and len(left) + len(right) == len(blocks):
        return sorted(left, key=_block_position_key) + sorted(right, key=_block_position_key)
    return sorted(blocks, key=_block_position_key)


def _block_position_key(block: LayoutBlock) -> tuple[float, float]:
    if block.bbox is None:
        return (float("inf"), float("inf"))
    return (round(block.bbox[1], 1), block.bbox[0])


def _classify_pdf_layout(
    pages: list[tuple[int, float, float, list[ParsedBlock]]],
) -> list[ParsedPage]:
    recurring: Counter[tuple[str, str]] = Counter()
    for _, _, height, blocks in pages:
        seen_on_page: set[tuple[str, str]] = set()
        for block in blocks:
            if block.block_type not in {"paragraph", "list"}:
                continue
            zone = _page_margin_zone(block, page_height=height)
            if zone is not None:
                seen_on_page.add((zone, _recurring_text_key(block.text)))
        recurring.update(seen_on_page)

    content_font_sizes = [
        block.font_size
        for _, _, height, blocks in pages
        for block in blocks
        if block.font_size is not None
        and len(block.text) >= 20
        and _page_margin_zone(block, page_height=height) is None
    ]
    if not content_font_sizes:
        content_font_sizes = [
            block.font_size
            for _, _, _, blocks in pages
            for block in blocks
            if block.font_size is not None
        ]
    body_size = float(median(content_font_sizes)) if content_font_sizes else 10.0

    preliminary: list[tuple[int, float, float, list[ParsedBlock]]] = []
    heading_sizes: set[float] = set()
    for page_no, width, height, blocks in pages:
        page_values: list[ParsedBlock] = []
        for block in blocks:
            if block.block_type not in {"paragraph", "list"}:
                page_values.append(block)
                continue
            zone = _page_margin_zone(block, page_height=height)
            recurring_margin = (
                zone is not None
                and len(pages) >= 2
                and recurring[(zone, _recurring_text_key(block.text))] >= 2
            )
            if zone == "header" and recurring_margin:
                classified = replace(block, block_type="header")
            elif zone == "footer" and (
                recurring_margin or _looks_like_page_number(block.text)
            ):
                classified = replace(block, block_type="footer")
            elif _is_heading_candidate(block, body_size=body_size):
                heading_sizes.add(round(block.font_size or body_size, 1))
                classified = replace(block, block_type="heading")
            elif _looks_like_list_item(block.text):
                classified = replace(block, block_type="list")
            else:
                classified = replace(block, block_type="paragraph")
            page_values.append(classified)
        preliminary.append((page_no, width, height, page_values))

    size_levels = {
        size: min(index + 1, 6)
        for index, size in enumerate(sorted(heading_sizes, reverse=True))
    }
    classified_pages: list[ParsedPage] = []
    for page_no, _, _, blocks in preliminary:
        classified_pages.append(
            ParsedPage(
                page_no=page_no,
                blocks=tuple(
                    replace(
                        block,
                        heading_level=_heading_level_from_text(block.text)
                        or size_levels.get(round(block.font_size or body_size, 1), 1),
                    )
                    if block.block_type == "heading"
                    else block
                    for block in blocks
                ),
            )
        )
    return classified_pages


def _page_margin_zone(block: LayoutBlock, *, page_height: float) -> str | None:
    if block.bbox is None or page_height <= 0:
        return None
    if block.bbox[3] <= page_height * 0.10:
        return "header"
    if block.bbox[1] >= page_height * 0.90:
        return "footer"
    return None


def _recurring_text_key(text: str) -> str:
    normalized = re.sub(r"\d+", "#", normalize_text(text).casefold())
    return re.sub(r"\s+", "", normalized)[:160]


def _looks_like_page_number(text: str) -> bool:
    normalized = normalize_text(text)
    return bool(
        re.fullmatch(r"(?:page\s*)?\d+(?:\s*(?:/|of)\s*\d+)?", normalized, re.IGNORECASE)
        or re.fullmatch(r"第?\s*\d+\s*页", normalized)
    )


def _looks_like_list_item(text: str) -> bool:
    return bool(
        re.match(
            r"^\s*(?:[-*•●▪◦]|[（(]?[0-9A-Za-z一二三四五六七八九十]+[）).、．])\s*",
            text,
        )
    )


def _is_heading_candidate(block: LayoutBlock, *, body_size: float) -> bool:
    text = normalize_text(block.text)
    if not text or len(text) > 160 or ("\n" in text and len(text.splitlines()) > 2):
        return False
    if _heading_level_from_text(text) is not None:
        return True
    if text.endswith(("。", "；", ";", "，", ",")) and len(text) > 25:
        return False
    font_size = block.font_size or body_size
    return font_size >= body_size * 1.16 or (block.is_bold and font_size >= body_size * 1.02)


def _heading_level_from_text(text: str) -> int | None:
    normalized = normalize_text(text)
    decimal = re.match(r"^(\d+(?:\.\d+){0,5})(?:[\s、．.)]|$)", normalized)
    if decimal:
        return min(decimal.group(1).count(".") + 1, 6)
    if re.match(r"^第[一二三四五六七八九十百0-9]+章", normalized):
        return 1
    if re.match(r"^第[一二三四五六七八九十百0-9]+节", normalized):
        return 2
    if re.match(r"^[一二三四五六七八九十百]+、", normalized):
        return 1
    if re.match(r"^[（(][一二三四五六七八九十百0-9]+[）)]", normalized):
        return 2
    if re.match(r"^(?:chapter|part)\s+[0-9ivxlcdm]+", normalized, re.IGNORECASE):
        return 1
    return None


def _reconstruct_pdf_sections(
    pages: list[ParsedPage],
) -> tuple[list[ParsedPage], tuple[Section, ...]]:
    roots: list[_SectionBuilder] = []
    stack: list[_SectionBuilder] = []
    rebuilt_pages: list[ParsedPage] = []
    for page in pages:
        rebuilt: list[LayoutBlock] = []
        for block in page.blocks:
            if block.block_type == "heading":
                level = block.heading_level or 1
                while stack and stack[-1].level >= level:
                    stack.pop()
                parent = stack[-1] if stack else None
                path = (*parent.path, block.text) if parent else (block.text,)
                current = _SectionBuilder(
                    title=block.text,
                    level=level,
                    path=path,
                    parent_path=parent.path if parent else None,
                )
                if parent is None:
                    roots.append(current)
                else:
                    parent.children.append(current)
                stack.append(current)
            section_path = stack[-1].path if stack else ()
            if page.page_no is not None:
                for section in stack:
                    section.pages.add(page.page_no)
            rebuilt.append(replace(block, section_path=section_path))
        rebuilt_pages.append(ParsedPage(page_no=page.page_no, blocks=tuple(rebuilt)))
    return rebuilt_pages, tuple(_freeze_section(item) for item in roots)


def _freeze_section(section: _SectionBuilder) -> Section:
    page_range = (min(section.pages), max(section.pages)) if section.pages else None
    return Section(
        title=section.title,
        level=section.level,
        parent_section=section.parent_path,
        children=tuple(_freeze_section(item) for item in section.children),
        page_range=page_range,
    )


def parse_docx(
    path: str | Path,
    *,
    ocr_enabled: bool = False,
    ocr_min_text_chars: int = 40,
    ocr_provider: OcrProvider | None = None,
    image_output_dir: str | Path | None = None,
    image_description_provider: ImageDescriptionProvider | None = None,
) -> ParsedDocument:
    resolved = Path(path).resolve(strict=True)
    return _parse_docx_document(
        document=DocxDocument(str(resolved)),
        source_path=str(resolved),
        ocr_enabled=ocr_enabled,
        ocr_min_text_chars=ocr_min_text_chars,
        ocr_provider=ocr_provider,
        image_output_dir=Path(image_output_dir) if image_output_dir is not None else None,
        image_description_provider=image_description_provider,
    )


def _parse_docx_document(
    *,
    document,
    source_path: str,
    ocr_enabled: bool,
    ocr_min_text_chars: int,
    ocr_provider: OcrProvider | None,
    image_output_dir: Path | None,
    image_description_provider: ImageDescriptionProvider | None,
) -> ParsedDocument:
    blocks: list[ParsedBlock] = []
    section_stack: list[str] = []

    image_ordinal = 0
    provider = ocr_provider or RapidOcrProvider() if ocr_enabled else None
    for paragraph in document.paragraphs:
        text = normalize_text(paragraph.text)
        if text:
            style_name = paragraph.style.name if paragraph.style is not None else ""
            heading_level = _heading_level(style_name)
            if heading_level is not None:
                section_stack = section_stack[: heading_level - 1]
                section_stack.append(text)
                block_type = "heading"
            else:
                block_type = "paragraph"
            blocks.append(
                ParsedBlock(
                    text=text,
                    page_no=None,
                    section_path=tuple(section_stack),
                    block_type=block_type,
                )
            )
        image_parts = _docx_paragraph_images(paragraph=paragraph, document=document)
        for relation_id, image_part in image_parts:
            image_ordinal += 1
            blocks.extend(
                _docx_image_blocks(
                    image_bytes=image_part.blob,
                    source_path=source_path,
                    relation_id=relation_id,
                    ordinal=image_ordinal,
                    section_path=tuple(section_stack),
                    anchor_text=text or "Image-only paragraph",
                    output_dir=image_output_dir,
                    provider=provider,
                    description_provider=image_description_provider,
                )
            )

    for table in document.tables:
        rows: list[str] = []
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells]
            if any(cells):
                rows.append("\t".join(cells))
        table_text = "\n".join(rows)
        if table_text:
            blocks.append(
                ParsedBlock(
                    text=table_text,
                    page_no=None,
                    section_path=tuple(section_stack),
                    block_type="table",
                )
            )
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for relation_id, image_part in _docx_paragraph_images(
                        paragraph=paragraph,
                        document=document,
                    ):
                        image_ordinal += 1
                        blocks.extend(
                            _docx_image_blocks(
                                image_bytes=image_part.blob,
                                source_path=source_path,
                                relation_id=relation_id,
                                ordinal=image_ordinal,
                                section_path=tuple(section_stack),
                                anchor_text=normalize_text(cell.text) or "Image in table cell",
                                output_dir=image_output_dir,
                                provider=provider,
                                description_provider=image_description_provider,
                            )
                        )

    for section in document.sections:
        for location, container in (("Header", section.header), ("Footer", section.footer)):
            for paragraph in container.paragraphs:
                for relation_id, image_part in _docx_paragraph_images(
                    paragraph=paragraph,
                    document=document,
                    part=container.part,
                ):
                    image_ordinal += 1
                    blocks.extend(
                        _docx_image_blocks(
                            image_bytes=image_part.blob,
                            source_path=source_path,
                            relation_id=relation_id,
                            ordinal=image_ordinal,
                            section_path=tuple(section_stack),
                            anchor_text=f"{location}: {normalize_text(paragraph.text)}",
                            output_dir=image_output_dir,
                            provider=provider,
                            description_provider=image_description_provider,
                        )
                    )

    return ParsedDocument(
        source_path=source_path,
        parser_version=PARSER_VERSION,
        pages=(ParsedPage(page_no=None, blocks=tuple(blocks)),),
    )


def _docx_paragraph_images(*, paragraph, document, part=None) -> list[tuple[str, Any]]:
    images: list[tuple[str, Any]] = []
    seen_relation_ids: set[str] = set()
    for blip in paragraph._p.xpath(".//*[local-name()='blip']"):
        relation_id = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
        if not relation_id or relation_id in seen_relation_ids:
            continue
        relation = (part or document.part).rels.get(relation_id)
        if relation is None or not relation.reltype.endswith("/image"):
            continue
        seen_relation_ids.add(relation_id)
        images.append((relation_id, relation.target_part))
    return images


def _docx_image_blocks(
    *,
    image_bytes: bytes,
    source_path: str,
    relation_id: str,
    ordinal: int,
    section_path: tuple[str, ...],
    anchor_text: str,
    output_dir: Path | None,
    provider: OcrProvider | None,
    description_provider: ImageDescriptionProvider | None,
) -> list[ParsedBlock]:
    image_path = _save_pdf_image(
        image_bytes=image_bytes,
        source_path=source_path,
        page_no=ordinal,
        output_dir=output_dir,
    )
    image_hash = hashlib.sha256(image_bytes).hexdigest()
    metadata = {
        "image_hash": image_hash,
        "ordinal": ordinal,
        "anchor": anchor_text[:240],
        "relation_id": relation_id,
        "path": str(image_path),
    }
    ocr_cells = []
    if provider is not None:
        try:
            ocr_cells = [
                _ocr_layout_block(item, page_no=0)
                for item in provider.extract_image(image_bytes=image_bytes)
                if normalize_text(item.text)
            ]
        except (OSError, RuntimeError, TypeError, ValueError):
            ocr_cells = []
    ocr_text = "\n".join(cell.text for cell in ocr_cells)
    if _looks_like_textual_image(ocr_cells):
        region = _union_bboxes(
            cell.bbox for cell in ocr_cells if cell.bbox is not None
        )
        tables = (
            _reconstruct_ocr_tables(
                region=region,
                page_no=0,
                context_blocks=ocr_cells,
                image_path=str(image_path),
            )
            if region is not None
            else []
        )
        if tables:
            return [
                replace(
                    table,
                    page_no=None,
                    source_type="docx_ocr_table",
                    caption="OCR reconstructed image table",
                    image_metadata={**metadata, "mode": "ocr_table"},
                )
                for table in tables
            ]
        caption = f"Image text:\n{ocr_text}"
        return [
            ParsedBlock(
                text=caption,
                page_no=None,
                section_path=section_path,
                block_type="image",
                source_type="docx_image_ocr",
                image_path=str(image_path),
                caption=caption,
                image_metadata={**metadata, "mode": "ocr"},
            )
        ]
    context = anchor_text or "No adjacent text is available."
    caption = f"Contextual image summary: {context}"
    if description_provider is not None:
        # No vision model is configured by default. A configured provider may
        # replace this context-only fallback with a grounded visual summary.
        try:
            described = normalize_text(
                description_provider.describe_image(
                    image_bytes=image_bytes,
                    page_no=0,
                    bbox=(0.0, 0.0, 1.0, 1.0),
                    nearby_text=context,
                    ocr_text="",
                )
            )
            if described:
                caption = described
        except (OSError, RuntimeError, TypeError, ValueError):
            pass
    return [
        ParsedBlock(
            text=caption,
            page_no=None,
            section_path=section_path,
            block_type="image",
            source_type="docx_image_context",
            image_path=str(image_path),
            caption=caption,
            image_metadata={**metadata, "mode": "context"},
        )
    ]


def _looks_like_textual_image(cells: list[ParsedBlock]) -> bool:
    text_length = sum(len(cell.text) for cell in cells)
    return len(cells) >= 2 and text_length >= 12


def parse_doc(
    path: str | Path,
    *,
    ocr_enabled: bool = False,
    ocr_min_text_chars: int = 40,
    ocr_provider: OcrProvider | None = None,
    image_output_dir: str | Path | None = None,
    image_description_provider: ImageDescriptionProvider | None = None,
) -> ParsedDocument:
    """Extract legacy Word content through the locally installed Microsoft Word."""
    resolved = Path(path).resolve(strict=True)
    pythoncom, client = _office_com_modules("Microsoft Word")
    application = document = None
    try:
        application = client.DispatchEx("Word.Application")
        application.Visible = False
        application.DisplayAlerts = 0
        document = application.Documents.Open(
            str(resolved),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
            Visible=False,
            OpenAndRepair=True,
        )
        with tempfile.TemporaryDirectory(prefix="enterprise-rag-doc-") as directory:
            converted = Path(directory) / "converted.docx"
            # Word preserves inline and floating pictures when converting to
            # DOCX, allowing the DOCX image pipeline to handle both formats.
            document.SaveAs2(str(converted), FileFormat=16, AddToRecentFiles=False)
            parsed = parse_docx(
                converted,
                ocr_enabled=ocr_enabled,
                ocr_min_text_chars=ocr_min_text_chars,
                ocr_provider=ocr_provider,
                image_output_dir=image_output_dir,
                image_description_provider=image_description_provider,
            )
        return replace(parsed, source_path=str(resolved))
    except Exception as exc:
        raise OfficeParsingError(f"unable to parse legacy Word document: {resolved.name}") from exc
    finally:
        if document is not None:
            document.Close(False)
        if application is not None:
            application.Quit()
        pythoncom.CoUninitialize()


def parse_presentation(path: str | Path) -> ParsedDocument:
    """Extract presentation content without relying on interactive Office automation."""
    resolved = Path(path).resolve(strict=True)
    if resolved.suffix.lower() == ".pptx":
        return parse_pptx(resolved)
    return _parse_legacy_presentation(resolved)


def parse_pptx(path: str | Path) -> ParsedDocument:
    """Read PPTX slide text directly from its Open XML package without COM automation."""
    resolved = Path(path).resolve(strict=True)
    try:
        with ZipFile(resolved) as package:
            slide_names = sorted(
                (
                    name
                    for name in package.namelist()
                    if re.fullmatch(r"ppt/slides/slide[0-9]+\.xml", name)
                ),
                key=_pptx_slide_sort_key,
            )
            pages = []
            for slide_index, slide_name in enumerate(slide_names, start=1):
                root = ElementTree.fromstring(package.read(slide_name))
                text_parts = [
                    normalized
                    for node in root.iter("{http://schemas.openxmlformats.org/drawingml/2006/main}t")
                    if (normalized := normalize_text(node.text or ""))
                ]
                blocks = (
                    (ParsedBlock(text="\n".join(text_parts), page_no=slide_index),)
                    if text_parts
                    else ()
                )
                pages.append(ParsedPage(page_no=slide_index, blocks=blocks))
    except (BadZipFile, ElementTree.ParseError, KeyError, OSError) as exc:
        raise OfficeParsingError(f"unable to parse PPTX presentation: {resolved.name}") from exc
    if not pages:
        raise OfficeParsingError(f"PPTX contains no slides: {resolved.name}")
    return ParsedDocument(
        source_path=str(resolved), parser_version=PARSER_VERSION, pages=tuple(pages)
    )


def _pptx_slide_sort_key(slide_name: str) -> int:
    match = re.search(r"slide([0-9]+)\.xml$", slide_name)
    return int(match.group(1)) if match else 0


def _parse_legacy_presentation(resolved: Path) -> ParsedDocument:
    """Extract legacy PPT text boxes and tables through local PowerPoint."""
    pythoncom, client = _office_com_modules("Microsoft PowerPoint")
    application = presentation = None
    try:
        with tempfile.TemporaryDirectory(prefix="enterprise-rag-ppt-") as directory:
            temporary_path = Path(directory) / "presentation.ppt"
            shutil.copyfile(resolved, temporary_path)
            application = client.DispatchEx("PowerPoint.Application")
            presentation = application.Presentations.Open(str(temporary_path), 1, 0, 0)
            pages: list[ParsedPage] = []
            for slide_index in range(1, presentation.Slides.Count + 1):
                slide = presentation.Slides(slide_index)
                blocks: list[ParsedBlock] = []
                for shape_index in range(1, slide.Shapes.Count + 1):
                    shape = slide.Shapes(shape_index)
                    if shape.HasTable:
                        table = shape.Table
                        rows = []
                        for row_index in range(1, table.Rows.Count + 1):
                            cells = [
                                normalize_text(
                                    str(
                                        table.Cell(row_index, col_index)
                                        .Shape.TextFrame.TextRange.Text
                                    )
                                )
                                for col_index in range(1, table.Columns.Count + 1)
                            ]
                            if any(cells):
                                rows.append("\t".join(cells))
                        if rows:
                            blocks.append(
                                ParsedBlock(
                                    text="\n".join(rows),
                                    page_no=slide_index,
                                    block_type="table",
                                )
                            )
                        continue
                    if shape.HasTextFrame and shape.TextFrame.HasText:
                        text = normalize_text(str(shape.TextFrame.TextRange.Text))
                        if text:
                            blocks.append(ParsedBlock(text=text, page_no=slide_index))
                pages.append(ParsedPage(page_no=slide_index, blocks=tuple(blocks)))
        return ParsedDocument(
            source_path=str(resolved), parser_version=PARSER_VERSION, pages=tuple(pages)
        )
    except Exception as exc:
        raise OfficeParsingError(f"unable to parse presentation: {resolved.name}") from exc
    finally:
        if presentation is not None:
            presentation.Close()
        if application is not None:
            application.Quit()
        pythoncom.CoUninitialize()


def _office_com_modules(application_name: str):
    try:
        import pythoncom
        import win32com.client
    except ImportError as exc:
        raise OfficeParsingError(f"{application_name} integration is unavailable") from exc
    pythoncom.CoInitialize()
    return pythoncom, win32com.client


def parse_text_file(path: str | Path, *, markdown: bool) -> ParsedDocument:
    resolved = Path(path).resolve(strict=True)
    return _parse_text_content(
        text=resolved.read_text(encoding="utf-8"),
        source_path=str(resolved),
        markdown=markdown,
    )


def _parse_text_content(*, text: str, source_path: str, markdown: bool) -> ParsedDocument:
    blocks: list[ParsedBlock] = []
    section_stack: list[str] = []
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if not paragraph_lines:
            return
        text = normalize_text("\n".join(paragraph_lines))
        paragraph_lines.clear()
        if text:
            blocks.append(
                ParsedBlock(
                    text=text,
                    page_no=None,
                    section_path=tuple(section_stack),
                )
            )

    for raw_line in text.splitlines():
        if markdown:
            heading = re.match(r"^(#{1,6})\s+(.+)$", raw_line)
            if heading:
                flush_paragraph()
                level = len(heading.group(1))
                title = normalize_text(heading.group(2))
                section_stack = section_stack[: level - 1]
                section_stack.append(title)
                blocks.append(
                    ParsedBlock(
                        text=title,
                        page_no=None,
                        section_path=tuple(section_stack),
                        block_type="heading",
                    )
                )
                continue
        if raw_line.strip():
            paragraph_lines.append(raw_line)
        else:
            flush_paragraph()
    flush_paragraph()

    return ParsedDocument(
        source_path=source_path,
        parser_version=PARSER_VERSION,
        pages=(ParsedPage(page_no=None, blocks=tuple(blocks)),),
    )


def parse_xlsx(path: str | Path) -> ParsedDocument:
    """Parse an Excel workbook into row blocks with formula and cached value context."""
    resolved = Path(path).resolve(strict=True)
    formula_workbook = load_workbook(filename=resolved, read_only=True, data_only=False)
    value_workbook = load_workbook(filename=resolved, read_only=True, data_only=True)
    try:
        return _parse_xlsx_workbooks(
            formula_workbook=formula_workbook,
            value_workbook=value_workbook,
            source_path=str(resolved),
        )
    finally:
        formula_workbook.close()
        value_workbook.close()


def _parse_xlsx_workbooks(*, formula_workbook, value_workbook, source_path: str) -> ParsedDocument:
    pages = []
    value_sheets = {worksheet.title: worksheet for worksheet in value_workbook.worksheets}
    for formula_sheet in formula_workbook.worksheets:
        value_sheet = value_sheets.get(formula_sheet.title)
        if value_sheet is None:
            raise ValueError(f"cached-value worksheet is missing: {formula_sheet.title}")
        pages.append(
            ParsedPage(
                page_no=None,
                blocks=tuple(
                    _excel_row_blocks(
                        sheet_name=formula_sheet.title,
                        rows=_xlsx_rows(
                            formula_rows=formula_sheet.iter_rows(values_only=True),
                            value_rows=value_sheet.iter_rows(values_only=True),
                        ),
                    )
                ),
            )
        )
    return ParsedDocument(
        source_path=source_path,
        parser_version=PARSER_VERSION,
        pages=tuple(pages),
    )


def _xlsx_rows(*, formula_rows, value_rows) -> Iterable[tuple[int, list[Any]]]:
    for row_number, (formula_row, value_row) in enumerate(
        zip_longest(formula_rows, value_rows, fillvalue=()), start=1
    ):
        yield (
            row_number,
            [
                _formula_with_cached_value(formula_value, cached_value)
                for formula_value, cached_value in zip_longest(
                    formula_row, value_row, fillvalue=None
                )
            ],
        )


def _formula_with_cached_value(formula_value: Any, cached_value: Any) -> Any:
    if (
        isinstance(formula_value, str)
        and formula_value.startswith("=")
        and cached_value is not None
    ):
        return f"{formula_value}（计算值：{_display_excel_value(cached_value)}）"
    return formula_value


def parse_xls(path: str | Path) -> ParsedDocument:
    """Parse a legacy XLS workbook using xlrd when the source requires it."""
    import xlrd

    resolved = Path(path).resolve(strict=True)
    workbook = xlrd.open_workbook(str(resolved), on_demand=True)
    try:
        return _parse_xls_workbook(workbook=workbook, source_path=str(resolved))
    finally:
        workbook.release_resources()


def _parse_xls_workbook(*, workbook, source_path: str) -> ParsedDocument:
    pages = []
    for sheet_name in workbook.sheet_names():
        sheet = workbook.sheet_by_name(sheet_name)
        rows = (
            (row_index, sheet.row_values(row_index - 1))
            for row_index in range(1, sheet.nrows + 1)
        )
        pages.append(
            ParsedPage(
                page_no=None,
                blocks=tuple(_excel_row_blocks(sheet_name=sheet_name, rows=rows)),
            )
        )
    return ParsedDocument(
        source_path=source_path,
        parser_version=PARSER_VERSION,
        pages=tuple(pages),
    )


def parse_tar_or_gzip(
    path: str | Path,
    *,
    ocr_enabled: bool,
    ocr_min_text_chars_per_page: int,
    ocr_render_dpi: int,
    ocr_provider: OcrProvider | None,
    max_members: int,
    max_member_bytes: int,
    max_uncompressed_bytes: int,
    max_compression_ratio: int,
) -> ParsedDocument:
    resolved = Path(path).resolve(strict=True)
    if tarfile.is_tarfile(resolved):
        return _parse_tar_archive(
            resolved,
            ocr_enabled=ocr_enabled,
            ocr_min_text_chars_per_page=ocr_min_text_chars_per_page,
            ocr_render_dpi=ocr_render_dpi,
            ocr_provider=ocr_provider,
            max_members=max_members,
            max_member_bytes=max_member_bytes,
            max_uncompressed_bytes=max_uncompressed_bytes,
            max_compression_ratio=max_compression_ratio,
        )

    member_path = resolved.with_suffix("").name
    if not _is_supported_archive_member(member_path):
        raise ValueError("gzip file does not contain a supported document type")
    with gzip.open(resolved, "rb") as stream:
        contents = _read_limited(stream, max_member_bytes)
    if len(contents) > max_uncompressed_bytes:
        raise ArchiveLimitError("archive uncompressed size limit exceeded")
    return _archive_document_from_pages(
        resolved,
        _parse_archive_member_pages(
            member_path=member_path,
            contents=contents,
            ocr_enabled=ocr_enabled,
            ocr_min_text_chars_per_page=ocr_min_text_chars_per_page,
            ocr_render_dpi=ocr_render_dpi,
            ocr_provider=ocr_provider,
        ),
    )


def _parse_tar_archive(
    resolved: Path,
    *,
    ocr_enabled: bool,
    ocr_min_text_chars_per_page: int,
    ocr_render_dpi: int,
    ocr_provider: OcrProvider | None,
    max_members: int,
    max_member_bytes: int,
    max_uncompressed_bytes: int,
    max_compression_ratio: int,
) -> ParsedDocument:
    pages: list[ParsedPage] = []
    total_uncompressed_bytes = 0
    archive_size = max(resolved.stat().st_size, 1)
    with tarfile.open(resolved, mode="r:*") as archive:
        members = [item for item in archive.getmembers() if item.isfile()]
        if len(members) > max_members:
            raise ArchiveLimitError(f"archive contains more than {max_members} files")
        for member in members:
            if not _is_supported_archive_member(member.name):
                continue
            if member.size > max_member_bytes:
                raise ArchiveLimitError(
                    f"archive member exceeds {max_member_bytes} bytes: {member.name}"
                )
            total_uncompressed_bytes += member.size
            _validate_archive_size_limits(
                total_uncompressed_bytes=total_uncompressed_bytes,
                archive_size=archive_size,
                max_uncompressed_bytes=max_uncompressed_bytes,
                max_compression_ratio=max_compression_ratio,
            )
            stream = archive.extractfile(member)
            if stream is None:
                continue
            with stream:
                contents = _read_limited(stream, max_member_bytes)
            pages.extend(
                _parse_archive_member_pages(
                    member_path=member.name,
                    contents=contents,
                    ocr_enabled=ocr_enabled,
                    ocr_min_text_chars_per_page=ocr_min_text_chars_per_page,
                    ocr_render_dpi=ocr_render_dpi,
                    ocr_provider=ocr_provider,
                )
            )
    return _archive_document_from_pages(resolved, pages)


def parse_external_archive(
    path: str | Path,
    *,
    ocr_enabled: bool,
    ocr_min_text_chars_per_page: int,
    ocr_render_dpi: int,
    ocr_provider: OcrProvider | None,
    max_members: int,
    max_member_bytes: int,
    max_uncompressed_bytes: int,
    max_compression_ratio: int,
) -> ParsedDocument:
    """Read RAR and 7z members through the locally available libarchive tool."""
    resolved = Path(path).resolve(strict=True)
    bsdtar = _find_bsdtar()
    try:
        listed = subprocess.run(
            [bsdtar, "-tf", str(resolved)],
            check=True,
            capture_output=True,
        )
    except (OSError, subprocess.CalledProcessError) as exc:
        raise ValueError(f"unable to open archive: {resolved.name}") from exc

    member_names = _decode_external_archive_names(listed.stdout).splitlines()
    members = [line.strip() for line in member_names if line.strip()]
    if len(members) > max_members:
        raise ArchiveLimitError(f"archive contains more than {max_members} files")

    pages: list[ParsedPage] = []
    total_uncompressed_bytes = 0
    archive_size = max(resolved.stat().st_size, 1)
    for member_path in members:
        if not _is_supported_archive_member(member_path):
            continue
        contents = _read_external_archive_member(bsdtar, resolved, member_path, max_member_bytes)
        total_uncompressed_bytes += len(contents)
        _validate_archive_size_limits(
            total_uncompressed_bytes=total_uncompressed_bytes,
            archive_size=archive_size,
            max_uncompressed_bytes=max_uncompressed_bytes,
            max_compression_ratio=max_compression_ratio,
        )
        pages.extend(
            _parse_archive_member_pages(
                member_path=member_path,
                contents=contents,
                ocr_enabled=ocr_enabled,
                ocr_min_text_chars_per_page=ocr_min_text_chars_per_page,
                ocr_render_dpi=ocr_render_dpi,
                ocr_provider=ocr_provider,
            )
        )
    return _archive_document_from_pages(resolved, pages)


def _find_bsdtar() -> str:
    configured = os.environ.get("BSDTAR_PATH")
    candidate = configured or shutil.which("bsdtar") or shutil.which("tar")
    if candidate is None:
        raise ValueError("RAR and 7z parsing requires a local bsdtar executable")
    return candidate


def _decode_external_archive_names(contents: bytes) -> str:
    for encoding in ("utf-8", "mbcs", "gb18030"):
        try:
            return contents.decode(encoding)
        except UnicodeDecodeError:
            continue
    return contents.decode("gb18030", errors="replace")


def _read_external_archive_member(
    bsdtar: str,
    archive_path: Path,
    member_path: str,
    max_member_bytes: int,
) -> bytes:
    process = subprocess.Popen(
        [bsdtar, "-xOf", str(archive_path), member_path],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    assert process.stdout is not None
    try:
        contents = _read_limited(process.stdout, max_member_bytes)
        _, stderr = process.communicate()
    except Exception:
        process.kill()
        process.communicate()
        raise
    if process.returncode != 0:
        raise ValueError(stderr.decode("utf-8", errors="replace").strip() or "archive read failed")
    return contents


def _read_limited(stream, max_bytes: int) -> bytes:
    contents = stream.read(max_bytes + 1)
    if len(contents) > max_bytes:
        raise ArchiveLimitError(f"archive member exceeds {max_bytes} bytes")
    return contents


def _validate_archive_size_limits(
    *,
    total_uncompressed_bytes: int,
    archive_size: int,
    max_uncompressed_bytes: int,
    max_compression_ratio: int,
) -> None:
    if total_uncompressed_bytes > max_uncompressed_bytes:
        raise ArchiveLimitError("archive uncompressed size limit exceeded")
    if total_uncompressed_bytes > archive_size * max_compression_ratio:
        raise ArchiveLimitError("archive compression ratio limit exceeded")


def _is_supported_archive_member(member_path: str) -> bool:
    normalized = member_path.replace("\\", "/")
    return (
        Path(normalized).suffix.lower() in ARCHIVE_MEMBER_SUFFIXES
        and ".." not in Path(normalized).parts
        and not normalized.startswith(("/", "__MACOSX/"))
    )


def _parse_archive_member_pages(
    *,
    member_path: str,
    contents: bytes,
    ocr_enabled: bool,
    ocr_min_text_chars_per_page: int,
    ocr_render_dpi: int,
    ocr_provider: OcrProvider | None,
) -> list[ParsedPage]:
    background_work_gate.wait_for_background_work()
    try:
        parsed = _parse_archive_member(
            member_path=member_path,
            contents=contents,
            ocr_enabled=ocr_enabled,
            ocr_min_text_chars_per_page=ocr_min_text_chars_per_page,
            ocr_render_dpi=ocr_render_dpi,
            ocr_provider=ocr_provider,
        )
    except (BadZipFile, OSError, RuntimeError, ValueError):
        return []
    return _prefix_archive_member(parsed=parsed, member_path=member_path)


def _archive_document_from_pages(resolved: Path, pages: list[ParsedPage]) -> ParsedDocument:
    if not pages:
        raise ValueError("archive contains no readable supported documents")
    return ParsedDocument(
        source_path=str(resolved), parser_version=PARSER_VERSION, pages=tuple(pages)
    )


def parse_zip(
    path: str | Path,
    *,
    ocr_enabled: bool,
    ocr_min_text_chars_per_page: int,
    ocr_render_dpi: int,
    ocr_provider: OcrProvider | None,
    max_members: int,
    max_member_bytes: int,
    max_uncompressed_bytes: int,
    max_compression_ratio: int,
) -> ParsedDocument:
    """Read supported archive members in memory without extracting them to disk."""
    resolved = Path(path).resolve(strict=True)
    pages: list[ParsedPage] = []
    total_uncompressed_bytes = 0
    try:
        with ZipFile(resolved) as archive:
            members = [item for item in archive.infolist() if not item.is_dir()]
            if len(members) > max_members:
                raise ArchiveLimitError(f"archive contains more than {max_members} files")
            for member in members:
                member_path = member.filename.replace("\\", "/")
                if (
                    not _is_supported_archive_member(member_path)
                    or member.flag_bits & 0x1
                ):
                    continue
                if member.file_size > max_member_bytes:
                    raise ArchiveLimitError(
                        f"archive member exceeds {max_member_bytes} bytes: {member_path}"
                    )
                total_uncompressed_bytes += member.file_size
                if total_uncompressed_bytes > max_uncompressed_bytes:
                    raise ArchiveLimitError("archive uncompressed size limit exceeded")
                if (
                    member.file_size > 0
                    and (
                        member.compress_size == 0
                        or member.file_size > member.compress_size * max_compression_ratio
                    )
                ):
                    raise ArchiveLimitError(
                        f"archive compression ratio limit exceeded: {member_path}"
                    )
                contents = archive.read(member)
                pages.extend(
                    _parse_archive_member_pages(
                        member_path=member_path,
                        contents=contents,
                        ocr_enabled=ocr_enabled,
                        ocr_min_text_chars_per_page=ocr_min_text_chars_per_page,
                        ocr_render_dpi=ocr_render_dpi,
                        ocr_provider=ocr_provider,
                    )
                )
    except BadZipFile as exc:
        raise ValueError("invalid ZIP archive") from exc
    if not pages:
        raise ValueError("archive contains no readable supported documents")
    return ParsedDocument(
        source_path=str(resolved),
        parser_version=PARSER_VERSION,
        pages=tuple(pages),
    )


def _parse_archive_member(
    *,
    member_path: str,
    contents: bytes,
    ocr_enabled: bool,
    ocr_min_text_chars_per_page: int,
    ocr_render_dpi: int,
    ocr_provider: OcrProvider | None,
) -> ParsedDocument:
    suffix = Path(member_path).suffix.lower()
    source_path = f"archive://{member_path}"
    if suffix == ".pdf":
        with pymupdf.open(stream=contents, filetype="pdf") as document:
            return _parse_pdf_document(
                document=document,
                source_path=source_path,
                ocr_enabled=ocr_enabled,
                ocr_min_text_chars_per_page=ocr_min_text_chars_per_page,
                ocr_render_dpi=ocr_render_dpi,
                ocr_provider=ocr_provider,
            )
    if suffix == ".docx":
        return _parse_docx_document(
            document=DocxDocument(BytesIO(contents)),
            source_path=source_path,
            ocr_enabled=ocr_enabled,
            ocr_min_text_chars=ocr_min_text_chars_per_page,
            ocr_provider=ocr_provider,
        )
    if suffix == ".doc":
        return _parse_archive_office_member(
            member_path=member_path,
            contents=contents,
            parser=parse_doc,
        )
    if suffix in {".ppt", ".pptx"}:
        return _parse_archive_office_member(
            member_path=member_path,
            contents=contents,
            parser=parse_presentation,
        )
    if suffix in {".xlsx", ".xlsm"}:
        formula_workbook = load_workbook(
            filename=BytesIO(contents), read_only=True, data_only=False
        )
        value_workbook = load_workbook(filename=BytesIO(contents), read_only=True, data_only=True)
        try:
            return _parse_xlsx_workbooks(
                formula_workbook=formula_workbook,
                value_workbook=value_workbook,
                source_path=source_path,
            )
        finally:
            formula_workbook.close()
            value_workbook.close()
    if suffix == ".xls":
        import xlrd

        workbook = xlrd.open_workbook(file_contents=contents, on_demand=True)
        try:
            return _parse_xls_workbook(workbook=workbook, source_path=source_path)
        finally:
            workbook.release_resources()
    return _parse_text_content(
        text=_decode_text_member(contents),
        source_path=source_path,
        markdown=suffix == ".md",
    )


def _parse_archive_office_member(*, member_path: str, contents: bytes, parser) -> ParsedDocument:
    source_path = f"archive://{member_path}"
    suffix = Path(member_path).suffix.lower()
    with tempfile.TemporaryDirectory(prefix="enterprise-rag-office-") as directory:
        member_file = Path(directory) / f"document{suffix}"
        member_file.write_bytes(contents)
        parsed = parser(member_file)
    return ParsedDocument(
        source_path=source_path,
        parser_version=PARSER_VERSION,
        pages=parsed.pages,
    )


def _prefix_archive_member(*, parsed: ParsedDocument, member_path: str) -> list[ParsedPage]:
    prefix = f"压缩包内文件：{member_path}"
    return [
        ParsedPage(
            page_no=page.page_no,
            blocks=tuple(
                replace(
                    block,
                    section_path=(prefix, *block.section_path),
                )
                for block in page.blocks
            ),
        )
        for page in parsed.pages
    ]


def _decode_text_member(contents: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return contents.decode(encoding)
        except UnicodeDecodeError:
            continue
    return contents.decode("utf-8", errors="replace")


def _excel_row_blocks(
    *,
    sheet_name: str,
    rows: Iterable[tuple[int, list[Any]]],
) -> list[ParsedBlock]:
    blocks: list[ParsedBlock] = []
    titles: list[str] = []
    headers: list[str] | None = None
    table_title = ""
    section_path = (f"工作表：{sheet_name}",)

    for row_number, values in rows:
        cells = [_display_excel_value(value) for value in values]
        populated = [(index, value) for index, value in enumerate(cells) if value]
        if not populated:
            continue
        if headers is None:
            if len(populated) < 2:
                titles.append(" ".join(value for _, value in populated))
                continue
            headers = _excel_headers(cells)
            table_title = titles[0] if titles else ""
            title = "\n".join(titles)
            if title:
                blocks.append(
                    ParsedBlock(
                        text=f"工作表：{sheet_name}\n{title}",
                        page_no=None,
                        section_path=section_path,
                        block_type="heading",
                    )
                )
            continue

        fields = [
            f"{headers[index]}：{value}"
            for index, value in populated
            if index < len(headers)
        ]
        if fields:
            context = f"表格标题：{table_title}\n" if table_title else ""
            blocks.append(
                ParsedBlock(
                    text=f"工作表：{sheet_name}\n{context}行：{row_number}\n"
                    + " | ".join(fields),
                    page_no=None,
                    section_path=section_path,
                    block_type="table_row",
                )
            )

    if headers is None and titles:
        blocks.append(
            ParsedBlock(
                text=f"工作表：{sheet_name}\n" + "\n".join(titles),
                page_no=None,
                section_path=section_path,
                block_type="paragraph",
            )
        )
    return blocks


def _excel_headers(values: list[str]) -> list[str]:
    seen: dict[str, int] = {}
    headers: list[str] = []
    for index, value in enumerate(values, start=1):
        base = value or f"列{index}"
        seen[base] = seen.get(base, 0) + 1
        headers.append(base if seen[base] == 1 else f"{base}（{seen[base]}）")
    return headers


def _display_excel_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, (datetime, date)):
        return value.isoformat(sep=" ") if isinstance(value, datetime) else value.isoformat()
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return normalize_text(str(value))


def _heading_level(style_name: str) -> int | None:
    match = re.match(r"Heading\s+([1-6])$", style_name)
    if not match:
        return None
    return int(match.group(1))


def _text_character_count(blocks: list[ParsedBlock]) -> int:
    return sum(sum(character.isalnum() for character in block.text) for block in blocks)


def _should_ocr_page(*, blocks: list[ParsedBlock], min_text_chars: int) -> bool:
    """OCR only pages with no meaningful text; text-rich PDF pages are skipped."""
    return _text_character_count(blocks) < min_text_chars
